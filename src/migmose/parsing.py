"""
contains functions for file handling and parsing.
"""

import re
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import DefaultDict, Generator, Union

import click
import docx
from docx.document import Document
from docx.oxml import CT_Tbl  # type:ignore[attr-defined]
from docx.table import Table, _Cell
from efoli import EdifactFormat, EdifactFormatVersion
from loguru import logger

from migmose.mig.segmentlayout import SegmentLayout, SegmentLayoutLine


def find_file_to_format(
    message_formats: list[EdifactFormat], edi_energy_repo: Path, format_version: EdifactFormatVersion
) -> dict[EdifactFormat, Path]:
    """
    finds the file with the message type in the input directory
    """
    input_dir = edi_energy_repo / Path("edi_energy_de") / Path(format_version)
    all_file_dict: DefaultDict[EdifactFormat, list[Path]] = defaultdict(list)
    file_dict: dict[EdifactFormat, Path] = {}
    for message_format in message_formats:
        for file in input_dir.iterdir():
            if "MIG" not in file.name or file.suffix != ".docx":
                continue
            if message_format is EdifactFormat.UTILMDG and "Gas" in file.name:
                all_file_dict[EdifactFormat.UTILMDG].append(file)
            elif message_format is EdifactFormat.UTILMDS and "Strom" in file.name:
                all_file_dict[EdifactFormat.UTILMDS].append(file)
            elif message_format in file.name:
                all_file_dict[message_format].append(file)
        if len(all_file_dict[message_format]) == 0:
            logger.warning(f"⚠️ No file found for {message_format}", fg="red")
            continue
        file_dict[message_format] = get_latest_file(all_file_dict[message_format])
    if file_dict:
        return file_dict
    logger.error("❌ No files found in the input directory.", fg="red")
    raise click.Abort()


def _extract_date(file_path: Path) -> tuple[datetime, Path]:
    # Regex to extract the date format YYYYMMDD from the filename as a string
    match = re.search(r"(\d{8})\.docx$", file_path.name)
    if match:
        # Return the date as a datetime object for comparison and the path for use
        return datetime.strptime(match.group(1), "%Y%m%d"), file_path
    logger.warning(
        f"⚠️ No timestamp in filename found in {file_path}."
        + "in case of multiple docx files in this path, it must be a "
        + "timestamp with format 'yyyyMMdd.docx' in filename.",
        fg="red",
    )
    raise click.Abort()


def get_latest_file(file_list: list[Path]) -> Path:
    """
    This function takes a list of docx files Path
    and returns the Path of the latest MIG docx file based on the timestamp in its name.
    The timestamp is assumed to be formatted as YYYYMMDD and located just before the ".docx" extension.

    Parameters:
        file_list (list of Path): A list containing file paths with timestamps.

    Returns:
        Path: The path of the latest file. Returns None if no valid date is found.
    """
    # Initialize variables to keep track of the latest file and date
    latest_file: Path
    latest_date: datetime | None = None

    for file_path in file_list:
        date, path = _extract_date(file_path)
        if latest_date is None or date > latest_date:
            latest_file = path
            latest_date = date

    # Return the path of the file with the latest date
    return latest_file


def get_paragraphs_up_to_diagram(parent: Union[Document, _Cell]) -> Generator[Table, None, None]:
    """Goes through paragraphs and tables"""
    # pylint: disable=protected-access
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Passed parent argument must be of type Document or _Cell")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_Tbl):
            yield Table(child, parent)


_row_regex = re.compile(r"^(?P<left>\t\d+\t)(?P<nr>\d{0,5})(?P<right>\t.*)$")
"""
https://regex101.com/r/vtF07B/2
"""


def _zfill_nr(row_str: str) -> str:
    match = _row_regex.match(row_str)
    if not match:
        return row_str
    left = match.group("left")
    nr = match.group("nr")
    right = match.group("right")
    return f"{left}{nr.zfill(5)}{right}"


def parse_raw_nachrichtenstrukturzeile(input_path: Path) -> tuple[list[str], dict[str, list[SegmentLayout]]]:
    """
    parses raw nachrichtenstrukturzeile from a table. returns list of raw lines
    """
    # pylint: disable=protected-access
    doc = docx.Document(str(input_path.absolute()))
    docx_objects = get_paragraphs_up_to_diagram(doc)
    mig_tables = []
    nachrichtenstruktur_header = "Status\tMaxWdh\n\tZähler\tNr\tBez\tSta\tBDEW\tSta\tBDEW\tEbene\tInhalt"
    segmentlayout_header = "\tStandard\tBDEW\n\tZähler\tNr\tBez\tSt\tMaxWdh\tSt\tMaxWdh\tEbene\tName"
    segmentlayout_tables = []
    for docx_object in docx_objects:
        segmentlayout_table = []
        for ind, line in enumerate(docx_object._cells):
            # marks the beginning of the complete nachrichtentruktur table
            if line.text == nachrichtenstruktur_header:
                mig_tables.extend([row.text for row in docx_object._cells[ind + 1 :]])
                break

        is_segmentlayout_table = docx_object._cells[0].text == segmentlayout_header
        if is_segmentlayout_table:
            for row in docx_object.rows:
                for cell in iter_visual_cells(row):
                    segmentlayout_table.append(cell)
        if segmentlayout_table:
            segmentlayout_tables.append(segmentlayout_table)
    # filter empty rows and headers
    mig_tables = [_zfill_nr(row) for row in mig_tables if row not in ("", "\n", nachrichtenstruktur_header)]
    segmentlayouts = process_segmentlayouts(segmentlayout_tables)
    return mig_tables, segmentlayouts


def iter_visual_cells(row: Table) -> Generator[_Cell, None, None]:
    """
    Iterate over visual cells in one row. Taken from https://github.com/python-openxml/python-docx/issues/344.
    """
    prior_tc = None
    for cell in row.cells:  # type:ignore [attr-defined]
        this_tc = cell._tc  # pylint: disable=protected-access
        if this_tc is prior_tc:  # skip cells pointing to same `<w:tc>` element
            continue
        yield cell
        prior_tc = this_tc


# pylint: disable=too-many-locals
def process_segmentlayouts(segmentlayout_tables: list[list[_Cell]]) -> dict[str, list[SegmentLayout]]:
    """
    Create Segmentlayouts from list of _Cell objects
    """
    segment_layouts_dict: defaultdict[str, list[SegmentLayout]] = defaultdict(list)
    for segmentlayout_table in segmentlayout_tables:
        start_index_annotations = 0
        start_index_example = 0
        ignore_cells = []
        for index, cell in enumerate(segmentlayout_table):
            cell_text = cell.text
            if cell_text == "Bez":
                ignore_cells.append(index)
            if cell_text == "Bemerkung:":
                ignore_cells.append(index)
                start_index_annotations = index
            if cell_text == "Beispiel:":
                start_index_example = index + 1
                break
        assert len(ignore_cells) != 0

        segment_zaehler = segmentlayout_table[ignore_cells[0] - 6].text.split("\t")[1]

        cell_index = []
        for index, header_ind in enumerate(ignore_cells[:-1]):
            cell_index.extend(list(range(header_ind + 7, ignore_cells[index + 1] - 6, 7)))
        unique_indentations = {
            segmentlayout_table[index].paragraphs[0].paragraph_format.left_indent
            for index in cell_index
            if segmentlayout_table[index].text != ""
        }
        indentations = {indentation: index for index, indentation in enumerate(sorted(unique_indentations))}
        raw_lines = []
        for i in cell_index:
            if segmentlayout_table[i].text != "":
                raw_lines.append(
                    SegmentLayoutLine(
                        bezeichnung=segmentlayout_table[i].text,
                        name=segmentlayout_table[i + 1].text,
                        standard_status=segmentlayout_table[i + 2].text,
                        standard_format=segmentlayout_table[i + 3].text,
                        bdew_status=segmentlayout_table[i + 4].text.replace("\t", ""),
                        bdew_format=segmentlayout_table[i + 5].text.replace("\t", ""),
                        anwendung=segmentlayout_table[i + 6].text,
                        indent=indentations[segmentlayout_table[i].paragraphs[0].paragraph_format.left_indent],
                    )
                )
            else:
                raw_lines[-1].bezeichnung += segmentlayout_table[i].text
                raw_lines[-1].name += segmentlayout_table[i + 1].text
                raw_lines[-1].standard_status += segmentlayout_table[i + 2].text
                raw_lines[-1].bdew_status += segmentlayout_table[i + 3].text
                raw_lines[-1].standard_format += segmentlayout_table[i + 4].text
                raw_lines[-1].bdew_format += segmentlayout_table[i + 5].text
                raw_lines[-1].anwendung += segmentlayout_table[i + 6].text
        segment_layouts_dict[segment_zaehler].append(
            SegmentLayout(
                struktur=raw_lines,
                bemerkung="".join(
                    cell.text for cell in segmentlayout_table[start_index_annotations : start_index_example - 1]
                ),
                beispiel="".join(cell.text for cell in segmentlayout_table[start_index_example:]),
            )
        )
    return segment_layouts_dict
